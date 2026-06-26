"""
Word (.docx) document parser.

Extracts: headings (with hierarchy), paragraphs, tables, and embedded images.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn

from src.models import SourceDocument
from src.parsing.source_document import _table_to_text
from src.parsing.source_document import _save_extracted_image


def parse_docx(path: Path) -> SourceDocument:
    """Parse a .docx file into a SourceDocument.

    Args:
        path: Absolute path to the .docx file.

    Returns:
        Fully populated SourceDocument.
    """
    doc = Document(str(path))

    title: str = path.stem
    sections: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    images: list[dict[str, Any]] = []
    full_text_parts: list[str] = []

    current_section: dict[str, Any] | None = None
    image_counter = 0
    table_counter = 0

    def _flush_section() -> None:
        nonlocal current_section
        if current_section and current_section.get("paragraphs"):
            sections.append(current_section)
        current_section = None

    def _ensure_section(heading: str = "", level: int = 0) -> dict[str, Any]:
        nonlocal current_section
        if current_section is None:
            current_section = {
                "heading": heading,
                "level": level,
                "paragraphs": [],
                "images": [],
                "tables": [],
            }
        return current_section

    # Extract document properties title if available
    if doc.core_properties.title:
        title = doc.core_properties.title

    # Iterate through the document body in order
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        # ── Paragraph ──────────────────────────────────────────────────
        if tag == "p":
            para = _find_paragraph(doc, element)
            if para is None:
                continue

            style = para.style.name if para.style else "Normal"
            text = para.text.strip()

            if not text:
                continue

            # Detect heading by style name
            is_heading = False
            heading_level = 0
            if style.startswith("Heading") or style.startswith("heading"):
                is_heading = True
                try:
                    heading_level = int(style.split()[-1])
                except ValueError:
                    heading_level = 1
            elif style.startswith("Title"):
                is_heading = True
                heading_level = 0
                if not title or title == path.stem:
                    title = text

            if is_heading:
                _flush_section()
                _ensure_section(text, heading_level)
                full_text_parts.append(f"\n{'#' * max(heading_level, 1)} {text}\n")
            else:
                sec = _ensure_section()
                sec["paragraphs"].append(text)
                full_text_parts.append(text)

            # Extract inline images from this paragraph
            for run in para.runs:
                for blip in run._element.findall(
                    qn("a:drawing") + "/" + qn("wp:inline") + "/" + qn("a:graphic")
                    + "/" + qn("a:graphicData") + "/" + qn("pic:pic") + "/" + qn("pic:blipFill")
                    + "/" + qn("a:blip")
                ):
                    embed = blip.get(qn("r:embed"))
                    if embed:
                        img_part = doc.part.related_parts.get(embed)
                        if img_part and hasattr(img_part, "blob"):
                            image_counter += 1
                            saved_path = _save_extracted_image(
                                img_part.blob, path.stem, image_counter
                            )
                            img_entry = {
                                "caption": text[:80] if text else f"Image {image_counter}",
                                "path_to_saved_image": saved_path,
                                "page": None,
                            }
                            images.append(img_entry)
                            if current_section is not None:
                                current_section.setdefault("images", []).append(img_entry)

        # ── Table ───────────────────────────────────────────────────────
        elif tag == "tbl":
            tbl = _find_table(doc, element)
            if tbl is None:
                continue

            table_counter += 1
            headers: list[str] = []
            rows: list[list[str]] = []

            for row_idx, row in enumerate(tbl.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if row_idx == 0:
                    headers = cells
                else:
                    rows.append(cells)

            table_entry = {
                "caption": f"Table {table_counter}",
                "headers": headers,
                "rows": rows,
                "page": None,
            }
            tables.append(table_entry)
            if current_section is not None:
                current_section.setdefault("tables", []).append(table_entry)
            # Add table text representation to full text
            full_text_parts.append(_table_to_text(headers, rows))

    _flush_section()

    # If no sections detected, create one with all text
    if not sections:
        sections.append({
            "heading": title,
            "level": 0,
            "paragraphs": [p for p in full_text_parts if p.strip()],
            "images": images,
            "tables": tables,
        })

    return SourceDocument(
        title=title,
        full_text="\n\n".join(full_text_parts),
        sections=sections,
        tables=tables,
        images=images,
        metadata={
            "filename": path.name,
            "type": "docx",
            "page_count": None,
            "section_count": len(sections),
            "table_count": len(tables),
            "image_count": len(images),
        },
    )


# ── Internal helpers ────────────────────────────────────────────────────────

def _find_paragraph(doc: Document, element: Any) -> Any:
    """Locate the python-docx Paragraph object matching an lxml element."""
    # Match by xml identity
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc: Document, element: Any) -> Any:
    """Locate the python-docx Table object matching an lxml element."""
    for tbl in doc.tables:
        if tbl._element is element:
            return tbl
    return None

from src.models import SourceDocument
from src.parsing.source_document import _table_to_text
from src.parsing.source_document import _save_extracted_image


def parse_docx(path: Path) -> SourceDocument:
    """Parse a .docx file into a SourceDocument.

    Args:
        path: Absolute path to the .docx file.

    Returns:
        Fully populated SourceDocument.
    """
    doc = Document(str(path))

    title: str = path.stem
    sections: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    images: list[dict[str, Any]] = []
    full_text_parts: list[str] = []

    current_section: dict[str, Any] | None = None
    image_counter = 0
    table_counter = 0

    def _flush_section() -> None:
        nonlocal current_section
        if current_section and current_section.get("paragraphs"):
            sections.append(current_section)
        current_section = None

    def _ensure_section(heading: str = "", level: int = 0) -> dict[str, Any]:
        nonlocal current_section
        if current_section is None:
            current_section = {
                "heading": heading,
                "level": level,
                "paragraphs": [],
                "images": [],
                "tables": [],
            }
        return current_section

    # Extract document properties title if available
    if doc.core_properties.title:
        title = doc.core_properties.title

    # Iterate through the document body in order
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        # ── Paragraph ──────────────────────────────────────────────────
        if tag == "p":
            para = _find_paragraph(doc, element)
            if para is None:
                continue

            style = para.style.name if para.style else "Normal"
            text = para.text.strip()

            if not text:
                continue

            # Detect heading by style name
            is_heading = False
            heading_level = 0
            if style.startswith("Heading") or style.startswith("heading"):
                is_heading = True
                try:
                    heading_level = int(style.split()[-1])
                except ValueError:
                    heading_level = 1
            elif style.startswith("Title"):
                is_heading = True
                heading_level = 0
                if not title or title == path.stem:
                    title = text

            if is_heading:
                _flush_section()
                _ensure_section(text, heading_level)
                full_text_parts.append(f"\n{'#' * max(heading_level, 1)} {text}\n")
            else:
                sec = _ensure_section()
                sec["paragraphs"].append(text)
                full_text_parts.append(text)

            # Extract inline images from this paragraph
            for run in para.runs:
                for blip in run._element.findall(
                    qn("a:drawing") + "/" + qn("wp:inline") + "/" + qn("a:graphic")
                    + "/" + qn("a:graphicData") + "/" + qn("pic:pic") + "/" + qn("pic:blipFill")
                    + "/" + qn("a:blip")
                ):
                    embed = blip.get(qn("r:embed"))
                    if embed:
                        img_part = doc.part.related_parts.get(embed)
                        if img_part and hasattr(img_part, "blob"):
                            image_counter += 1
                            saved_path = _save_extracted_image(
                                img_part.blob, path.stem, image_counter
                            )
                            img_entry = {
                                "caption": text[:80] if text else f"Image {image_counter}",
                                "path_to_saved_image": saved_path,
                                "page": None,
                            }
                            images.append(img_entry)
                            if current_section is not None:
                                current_section.setdefault("images", []).append(img_entry)

        # ── Table ───────────────────────────────────────────────────────
        elif tag == "tbl":
            tbl = _find_table(doc, element)
            if tbl is None:
                continue

            table_counter += 1
            headers: list[str] = []
            rows: list[list[str]] = []

            for row_idx, row in enumerate(tbl.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if row_idx == 0:
                    headers = cells
                else:
                    rows.append(cells)

            table_entry = {
                "caption": f"Table {table_counter}",
                "headers": headers,
                "rows": rows,
                "page": None,
            }
            tables.append(table_entry)
            if current_section is not None:
                current_section.setdefault("tables", []).append(table_entry)
            # Add table text representation to full text
            full_text_parts.append(_table_to_text(headers, rows))

    _flush_section()

    # If no sections detected, create one with all text
    if not sections:
        sections.append({
            "heading": title,
            "level": 0,
            "paragraphs": [p for p in full_text_parts if p.strip()],
            "images": images,
            "tables": tables,
        })

    return SourceDocument(
        title=title,
        full_text="\n\n".join(full_text_parts),
        sections=sections,
        tables=tables,
        images=images,
        metadata={
            "filename": path.name,
            "type": "docx",
            "page_count": None,
            "section_count": len(sections),
            "table_count": len(tables),
            "image_count": len(images),
        },
    )


# ── Internal helpers ────────────────────────────────────────────────────────

def _find_paragraph(doc: Document, element: Any) -> Any:
    """Locate the python-docx Paragraph object matching an lxml element."""
    # Match by xml identity
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc: Document, element: Any) -> Any:
    """Locate the python-docx Table object matching an lxml element."""
    for tbl in doc.tables:
        if tbl._element is element:
            return tbl
    return None


def _table_to_text(headers: list[str], rows: list[list[str]]) -> str:
    """Render a table as markdown-ish text for the full_text field."""
    lines = [" | ".join(headers)]
    lines.append(" | ".join(["---"] * len(headers)))
    for row in rows:
        # Pad row to match header length
        padded = row + [""] * (len(headers) - len(row))
        lines.append(" | ".join(padded[:len(headers)]))
    return "\n".join(lines)
